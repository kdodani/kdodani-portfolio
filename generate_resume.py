from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

# ── Portfolio color system (matches globals.css + design-system.md) ───────────
# Background: #faf9f6  (warm off-white — body bg)
# Headings:   #1c1917  (stone-900 — near-black warm)
# Body text:  #44403c  (stone-700)
# Muted:      #78716c  (stone-500)
# Accent:     #6D28D9  (--accent-hero-1 purple)
# Mid-accent: #5B21B6  (--accent-strong)
# Sub-accent: #4F46E5  (--accent-hero-2 indigo)
DARK    = RGBColor(0x1c, 0x19, 0x17)   # stone-900 — headings
MID     = RGBColor(0x5B, 0x21, 0xB6)   # accent-strong purple — section labels
ACCENT  = RGBColor(0x6D, 0x28, 0xD9)   # accent-hero-1 purple — rules / name accent
BODY    = RGBColor(0x44, 0x40, 0x3c)   # stone-700 — body text
LIGHT   = RGBColor(0x78, 0x71, 0x6c)   # stone-500 — muted / metadata
BG_HEX  = "FAF9F6"                     # warm off-white page background

def set_page_background(doc, hex_color):
    """Set the document page background color."""
    body_elem = doc.element.body
    doc_elem  = body_elem.getparent()
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), hex_color)
    doc_elem.insert(0, bg)
    settings = doc.settings.element
    disp = OxmlElement('w:displayBackgroundShape')
    settings.insert(0, disp)

set_page_background(doc, BG_HEX)

def clear_paragraph_space(p, before=0, after=0):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)

def add_horizontal_rule(doc, color=RGBColor(0x6D, 0x28, 0xD9), thickness=6):
    """Insert a thin colored paragraph border as a visual divider."""
    p = doc.add_paragraph()
    clear_paragraph_space(p, before=2, after=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading(doc, text, size=10.5, bold=True, color=RGBColor(0x6D, 0x28, 0xD9), space_before=8, space_after=2, caps=True):
    p = doc.add_paragraph()
    clear_paragraph_space(p, before=space_before, after=space_after)
    run = p.add_run(text.upper() if caps else text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def subheading(doc, role, org, period, size=10):
    """Role | Org   ·   Period  — two runs on one line."""
    p = doc.add_paragraph()
    clear_paragraph_space(p, before=5, after=1)
    r1 = p.add_run(f"{role}  |  {org}")
    r1.bold = True
    r1.font.size = Pt(size)
    r1.font.color.rgb = DARK
    r2 = p.add_run(f"    {period}")
    r2.bold = False
    r2.font.size = Pt(9)
    r2.font.color.rgb = LIGHT
    return p

def sublabel(doc, text, size=9, italic=True, color=RGBColor(0x4F, 0x46, 0xE5), space_before=0, space_after=1):
    p = doc.add_paragraph()
    clear_paragraph_space(p, before=space_before, after=space_after)
    run = p.add_run(text)
    run.italic = italic
    run.bold = not italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def bullet(doc, text, size=9.5, indent=Inches(0.2), hanging=Inches(0.2)):
    p = doc.add_paragraph(style='List Bullet')
    clear_paragraph_space(p, before=1, after=1)
    p.paragraph_format.left_indent   = indent + hanging
    p.paragraph_format.first_line_indent = -hanging
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = BODY
    return p

def body(doc, text, size=9.5, color=BODY, space_before=2, space_after=2, italic=False):
    p = doc.add_paragraph()
    clear_paragraph_space(p, before=space_before, after=space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    return p

# ══════════════════════════════════════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════════════════════════════════════
p_name = doc.add_paragraph()
clear_paragraph_space(p_name, before=0, after=2)
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_name.add_run("KHUSHBOO DODANI")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = DARK

p_contact = doc.add_paragraph()
clear_paragraph_space(p_contact, before=0, after=2)
p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
rc = p_contact.add_run(
    "San Francisco Bay Area  |  khushdodani@gmail.com  |  LinkedIn  |  GitHub  |  Portfolio"
)
rc.font.size = Pt(9)
rc.font.color.rgb = LIGHT

p_title = doc.add_paragraph()
clear_paragraph_space(p_title, before=2, after=2)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = p_title.add_run("PRODUCT MANAGER")
rt.bold = True
rt.font.size = Pt(12)
rt.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)   # accent-hero-1 purple

p_pills = doc.add_paragraph()
clear_paragraph_space(p_pills, before=1, after=4)
p_pills.alignment = WD_ALIGN_PARAGRAPH.CENTER
rp = p_pills.add_run(
    "Care Delivery Systems  •  EHR & Clinical Workflows  •  Clinician Experience  •  AI-enabled Operations"
)
rp.font.size = Pt(9)
rp.font.color.rgb = RGBColor(0x5B, 0x21, 0xB6)   # accent-strong purple

add_horizontal_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
body(
    doc,
    "Product Manager with 4+ years building care delivery systems, EHR-adjacent clinical workflows, and "
    "AI-enabled operations for healthcare platforms. Experienced partnering with engineers, clinicians, and "
    "clinical operations teams to reduce administrative friction, automate manual processes, and help clinicians "
    "focus on patients — in regulated healthcare environments where reliability, compliance, and trust are "
    "non-negotiable. Track record taking complex workflow challenges from discovery through measurable improvement "
    "with cross-functional teams.",
    size=9.5,
    space_before=4,
    space_after=4,
)

add_horizontal_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# EXPERIENCE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "Experience")

# ── WorkSafeBC ────────────────────────────────────────────────────────────────
subheading(doc, "Product Manager — Provider & Clinician Experience", "WorkSafeBC", "Apr 2024 – Present")

sublabel(doc, "Care Delivery Systems")
bullet(doc,
    "Own product strategy for digital systems supporting healthcare providers and clinicians submitting "
    "treatment requests, referrals, and invoices through WorkSafeBC's clinical portal.")
bullet(doc,
    "Partner with clinicians, clinical operations, engineering, architects, and executive leadership to "
    "improve the care delivery infrastructure supporting providers across British Columbia.")

sublabel(doc, "Clinical Workflow Infrastructure & EHR-Adjacent Systems")
bullet(doc,
    "Led redesign of clinician-facing workflows after interviews, usability studies, analytics, and journey "
    "mapping identified major friction in the clinical documentation and invoicing process.")
bullet(doc,
    "Reduced submission journey from 11 steps to 5, cutting clinician administrative burden and increasing "
    "digital adoption to 60% within six months — generating approximately $200K in annual operational savings.")
bullet(doc,
    "Improved data quality through front-end validation and simplified workflows, decreasing support calls "
    "by 30% while increasing clinician satisfaction (NPS +25).")
bullet(doc,
    "Established adoption metrics using PostHog and continuously prioritized improvements using clinician "
    "feedback and behavioral analytics.")

sublabel(doc, "Healthcare Interoperability & EHR Platform Strategy")
bullet(doc,
    "Led product discovery for a healthcare interoperability platform enabling direct submission of clinical "
    "documentation from EHR/EMR systems — including OSCAR — directly into WorkSafeBC, reducing documentation "
    "burden on clinicians and informing long-term platform strategy.")
bullet(doc,
    "Secured $80K in executive funding to expand provider connectivity through API-enabled integrations with "
    "the province's largest EHR vendors, establishing a roadmap covering approximately 70% of the provider ecosystem.")
bullet(doc,
    "Evaluated clinician workflows, EHR integration constraints, and technical tradeoffs to prioritize scalable "
    "interoperability capabilities across external healthcare systems.")

sublabel(doc, "Internal Systems & Clinical Access Infrastructure")
bullet(doc,
    "Led migration from a legacy government identity platform to a modern CIAM solution supporting secure "
    "access to referral and claims systems for 5,600 healthcare providers across 500+ organizations.")
bullet(doc,
    "Delivered migration with 0% downtime, 85% provider migration before deadline, and 92% post-launch login "
    "success — ensuring uninterrupted clinician access to critical care delivery systems.")
bullet(doc,
    "Coordinated Engineering, Operations, Communications, and external partners across rollout, provider "
    "onboarding, adoption, and change management.")

sublabel(doc, "AI-enabled Clinical Operations")
bullet(doc,
    "Led development of Referral Radar, an AI-assisted clinical decision support prototype using LLM agents, "
    "knowledge visualization, and intelligent task routing to recommend healthcare program referrals — "
    "demonstrating potential to reduce claim costs by approximately $20M.")
bullet(doc,
    "Leveraged AI-enabled prototyping tools including Lovable and GitHub Copilot to accelerate clinical product "
    "discovery, user validation, and engineering collaboration, reducing concept-to-validation timelines by "
    "approximately 40%.")

# ── Technical Safety BC ───────────────────────────────────────────────────────
subheading(doc, "Product Manager", "Technical Safety BC", "Oct 2021 – Apr 2024")
bullet(doc,
    "Led discovery for a mobile remote inspections platform through contextual user research, workflow mapping, "
    "usability testing, rapid prototyping, and experimentation before development investment.")
bullet(doc,
    "Partnered with engineering, field inspectors, designers, and operations leaders to translate operational "
    "pain points into product requirements, roadmap priorities, and implementation plans.")
bullet(doc,
    "Built NLP-powered email classification to automate routing and triage of incoming operational requests, "
    "reducing manual inbox processing burden — direct predecessor to AI-assisted workflow tooling.")
bullet(doc,
    "Built operational analytics products identifying an average 11.7% annual churn rate, enabling leadership "
    "to prioritize retention initiatives through predictive insights.")
bullet(doc,
    "Developed an RFM segmentation model driving targeted engagement campaigns that retained 78% of "
    "high-risk customers.")

# ── 180 Degrees ───────────────────────────────────────────────────────────────
subheading(doc, "Strategy Consultant", "180 Degrees Consulting — University of British Columbia", "Sep 2020 – Oct 2021")
bullet(doc,
    "Developed product and market strategies for nonprofit and clean energy organizations through customer "
    "research, market analysis, and strategic planning.")
bullet(doc,
    "Designed a data-driven acquisition strategy for Open Primaries that reduced website bounce rate by 37% "
    "while doubling subscriptions.")

# ── Convertus ─────────────────────────────────────────────────────────────────
subheading(doc, "SEO Consultant", "Convertus", "May 2018 – Jul 2019")
bullet(doc,
    "Built analytics-driven growth strategies for 75+ client websites, increasing organic performance by "
    "approximately 150%.")
bullet(doc,
    "Developed a forecasting tool predicting SEO performance with approximately 85% accuracy to improve "
    "marketing planning.")

add_horizontal_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# EDUCATION
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "Education")

subheading(doc, "Master of Business Analytics", "University of British Columbia", "")
body(doc, "Python  •  SQL  •  Machine Learning  •  Optimization  •  Product Analytics", size=9, color=LIGHT)

subheading(doc, "Bachelor of Business Administration (Marketing)", "British Columbia Institute of Technology", "")

add_horizontal_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# LEADERSHIP
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "Leadership")

sublabel(doc, "Speaker — ProductCamp BC (2025)", italic=False, size=9.5, color=DARK)
body(doc,
    '"How Legacy Industries Can Adopt a Product Mindset to Improve Service Delivery."',
    size=9.5, italic=True, space_before=0)

sublabel(doc, "Mentor — ProductBC", italic=False, size=9.5, color=DARK, space_before=4)
body(doc,
    "Mentor aspiring Product Managers in product discovery, strategy, and career development.",
    size=9.5, space_before=0)

add_horizontal_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SKILLS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "Skills")

def skills_row(doc, label, skills):
    p = doc.add_paragraph()
    clear_paragraph_space(p, before=2, after=2)
    rl = p.add_run(f"{label}:  ")
    rl.bold = True
    rl.font.size = Pt(9)
    rl.font.color.rgb = RGBColor(0x5B, 0x21, 0xB6)
    rs = p.add_run(skills)
    rs.font.size = Pt(9)
    rs.font.color.rgb = BODY

skills_row(doc, "Product",
    "Product Strategy  •  Product Discovery  •  Clinical Workflow Infrastructure  •  Care Delivery Systems  •  "
    "Internal Tools  •  Product Roadmaps  •  Workflow Automation  •  Platform Products  •  Product Analytics  •  "
    "Inbox Automation")
skills_row(doc, "Healthcare",
    "Clinician Experience  •  EHR Systems  •  Provider Experience  •  Healthcare Operations  •  "
    "Interoperability  •  Clinical Documentation  •  Telehealth  •  HIPAA  •  Clinical Systems")
skills_row(doc, "Technical",
    "REST APIs  •  EMR/EHR Integrations  •  CIAM/IAM  •  Platform Integrations  •  SQL  •  Python  •  "
    "Azure  •  Angular  •  .NET  •  PostHog  •  GitHub Copilot  •  Lovable")
skills_row(doc, "Research",
    "Journey Mapping  •  User Interviews  •  Workflow Analysis  •  Usability Testing  •  A/B Testing  •  "
    "Prototyping  •  Figma")
skills_row(doc, "AI",
    "Generative AI  •  LLM Applications  •  NLP  •  AI-assisted Product Development  •  Intelligent Task Routing")

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"c:\Users\khush\Desktop\kdodani-portfolio\Khushboo_Dodani_Resume_MidiHealth.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
